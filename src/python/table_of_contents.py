#!/usr/bin/env python

from bs4 import BeautifulSoup
from docx import Document
import os

ioc = '../../resources/table_of_contents.html'
with open(ioc) as fp:
    doc = BeautifulSoup(fp, features="lxml")

out_file = '../../resources/inhaltsverzeichnis.docx'
if os.path.isfile(out_file) :
  os.remove(out_file)
word_document = Document()

word_document.add_heading('Inhaltsverzeichnis', level=1)
for node in doc.body.children:
    if(node.name):
        if(node.string):
          entry = node.string.strip()
          word_document.add_paragraph(entry, style='List Number')
        else:
          entry = node.contents[0].strip()
          word_document.add_paragraph(entry, style='List Number')
          #if(len(node.contents) > 1):
          print('############### {}##############'.format(node.name))
          ordered_list = node.find_all("ol", limit = 1);
          [word_document.add_paragraph(line, style='List Number 2')
           for line in ordered_list.contents]



            #for li1 in node.contents[1].children:
            #  if(li1.name and li1.string):
            #    word_document.add_paragraph(li1.string.strip(), style='List Number 2')
word_document.save(out_file)
